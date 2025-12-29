"""
Generador de Memorias Descriptivas en Word con formato oficial para proyectos de mensura.
"""

import os
import logging
from datetime import datetime
from django.conf import settings
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)

class DocxGenerator:
    """Genera memoria descriptiva con formato oficial"""

    def __init__(self, plano):
        self.plano = plano
        self.datos = plano.datos_procesados or {}
        logger.debug("Datos recibidos en DocxGenerator: %s", self.datos)

    # -------------------------
    # Helpers
    # -------------------------
    def _add_styled_paragraph(self, doc, label, value):
        """Agrega un párrafo con etiqueta en negrita y valor normal"""
        p = doc.add_paragraph()
        run_label = p.add_run(f"{label}: ")
        run_label.bold = True
        p.add_run(str(value if value not in [None, ""] else "No especificado"))
        return p

    def _format_dominios(self):
        dominios = self.datos.get("dominios", [])
        if not dominios:
            return "No especificado"
        return ", ".join(
            [d.get("matricula", "") if isinstance(d, dict) else str(d) for d in dominios]
        ) or "No especificado"

    def _format_propietarios(self):
        propietarios = self.datos.get("propietarios", [])
        if not propietarios:
            return "No especificado"
        nombres = []
        for p in propietarios:
            if isinstance(p, dict):
                nom = p.get("nombre", "").strip()
                if nom:
                    nombres.append(nom)
            else:
                s = str(p).strip()
                if s:
                    nombres.append(s)
        if not nombres:
            return "No especificado"
        return ", ".join(nombres[:3]) + (" entre otros" if len(nombres) > 3 else "")

    def _dedupe_lados(self, lados):
        """Deduplica lados por (lado, medida) y mantiene el primero"""
        seen = set()
        result = []
        for lado_item in lados or []:
            if not isinstance(lado_item, dict):
                continue
            key = (
                str(lado_item.get("lado", "")).strip(),
                str(lado_item.get("mide", "")).strip()
            )
            if key in seen:
                continue
            seen.add(key)
            result.append(lado_item)
        return result

    def _normalize_refs(self, referencias):
        """Normaliza referencias a lista de strings sin duplicados"""
        if not referencias:
            return []
        if isinstance(referencias, str):
            refs = [r.strip() for r in referencias.split("\n") if r.strip()]
        elif isinstance(referencias, list):
            refs = [str(r).strip() for r in referencias if str(r).strip()]
        else:
            refs = [str(referencias).strip()]
        dedup, seen = [], set()
        for r in refs:
            if r not in seen:
                seen.add(r)
                dedup.append(r)
        return dedup

    # -------------------------
    # Tablas
    # -------------------------
    def _add_table_superficies(self, doc):
        h = doc.add_paragraph("3. PLANILLA DE SUPERFICIES")
        h.runs[0].bold = True

        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text, hdr_cells[1].text, hdr_cells[2].text = (
            "Designación",
            "Superficie",
            "Observaciones",
        )
        for cell in hdr_cells:
            if cell.paragraphs and cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].bold = True

        superficies = self.datos.get("superficies", [])
        if superficies:
            for sup in superficies:
                row_cells = table.add_row().cells
                if isinstance(sup, dict):
                    row_cells[0].text = sup.get("designacion", "No especificado")
                    st = sup.get("sup_titulo")
                    sm = sup.get("sup_mensura")
                    dif = sup.get("diferencia")
                    partes = [p for p in [st, sm, dif] if p and p != "No especificado"]
                    valor = " / ".join(partes) if partes else "No especificado"
                    row_cells[1].text = valor
                    row_cells[2].text = sup.get("observaciones", "") or " "
                else:
                    row_cells[0].text = "No especificado"
                    row_cells[1].text = str(sup)
                    row_cells[2].text = " "
        else:
            row_cells = table.add_row().cells
            row_cells[0].text = row_cells[1].text = row_cells[2].text = "No especificado"

        nota = "(Nota: Los valores exactos de cada superficie deben transcribirse de la Planilla de Superficies del plano)"
        doc.add_paragraph(nota)

    def _add_table_lados(self, doc):
        h = doc.add_paragraph("4. PLANILLA DE LADOS")
        h.runs[0].bold = True

        lados_raw = self.datos.get("lados", [])
        lados = self._dedupe_lados(lados_raw)

        table = doc.add_table(rows=1, cols=6)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text, hdr[4].text, hdr[5].text = (
            "Vértice",
            "Rumbo",
            "Lado",
            "Medida (m)",
            "Ángulo",
            "Linderos",
        )
        for cell in hdr:
            if cell.paragraphs and cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].bold = True

        if lados:
            for lado in lados:
                row = table.add_row().cells
                row[0].text = str(lado.get("vertice", "")) or "—"
                row[1].text = str(lado.get("rumbo", "")) or "—"
                row[2].text = str(lado.get("lado", "")) or "—"
                row[3].text = str(lado.get("mide", "")) or "—"
                row[4].text = str(lado.get("angulo", "")) or "—"
                row[5].text = str(lado.get("linderos", "")) or "—"
        else:
            row = table.add_row().cells
            for i in range(6):
                row[i].text = "No especificado"

    def _add_table_coordenadas(self, doc):
        coords = self.datos.get("coordenadas", [])
        if not coords:
            return
        h = doc.add_paragraph("7. COORDENADAS GEODÉSICAS")
        h.runs[0].bold = True
        table = doc.add_table(rows=1, cols=5)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text, hdr[4].text = (
            "Punto", "Latitud", "Longitud", "Norte GK", "Este GK"
        )
        for cell in hdr:
            if cell.paragraphs and cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].bold = True
        for c in coords:
            row = table.add_row().cells
            row[0].text = str(c.get("punto", "")) or "—"
            row[1].text = str(c.get("latitud", "")) or "—"
            row[2].text = str(c.get("longitud", "")) or "—"
            row[3].text = str(c.get("norte_gk", "")) or "—"
            row[4].text = str(c.get("este_gk", "")) or "—"

    # -------------------------
    # Generación principal
    # -------------------------
    def generate_memoria(self):
        try:
            logger.debug(
                "Generando memoria para plano %s con datos: %s",
                getattr(self.plano, "id", "sin_id"),
                self.datos,
            )

            doc = Document()

            # Estilo global
            style = doc.styles["Normal"]
            font = style.font
            font.name = "Times New Roman"
            font.size = Pt(12)

            # Título
            titulo = doc.add_paragraph()
            titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_t = titulo.add_run("MEMORIA DESCRIPTIVA")
            run_t.bold = True
            run_t.font.size = Pt(12)

            # Fecha de generación
            doc.add_paragraph(f"Documento generado el {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")

            # Cabecera
            self._add_styled_paragraph(doc, "DEPARTAMENTO", self.datos.get("departamento", "No especificado"))
            padrones = ", ".join(self.datos.get("padrones", [])).strip() or "No especificado"
            self._add_styled_paragraph(doc, "PADRON", padrones)
            self._add_styled_paragraph(doc, "LUGAR", self.datos.get("lugar", "No especificado"))
            self._add_styled_paragraph(doc, "DOMINIO", self._format_dominios())
            if self.datos.get("baricentro"):
                self._add_styled_paragraph(doc, "BARICENTRO GEOGRÁFICO", self.datos.get("baricentro"))
            self._add_styled_paragraph(doc, "OBJETO", self.datos.get("objeto", "No especificado"))
            self._add_styled_paragraph(doc, "INMUEBLE", self.datos.get("inmueble", "No especificado"))
            self._add_styled_paragraph(doc, "TITULAR", self._format_propietarios())
            fecha_op = self.datos.get("fecha_operaciones") or "No especificado"
            self._add_styled_paragraph(doc, "FECHA DE OPERACIÓN", fecha_op)

            # Sección 1: Extracto de Título
            h1 = doc.add_paragraph("1. EXTRACTO DE TÍTULO")
            h1.runs[0].bold = True
            doc.add_paragraph(f"Dominio: {self._format_dominios()}")
            doc.add_paragraph(f"Inmueble: {self.datos.get('inmueble', 'No especificado')}")
            medidas_linderos = self.datos.get("medidas_linderos") or "Según plano de mensura."
            doc.add_paragraph(f"Medidas y Linderos: {medidas_linderos}")

            # Sección 2: Descripción de las Operaciones
            h2 = doc.add_paragraph("2. DESCRIPCIÓN DE LAS OPERACIONES")
            h2.runs[0].bold = True
            descripcion = (self.datos.get("descripcion") or "").strip()
            if descripcion:
                doc.add_paragraph(descripcion)
            nota1 = (self.datos.get("nota1") or "").strip()
            nota2 = (self.datos.get("nota2") or "").strip()
            if nota1:
                doc.add_paragraph(f"Nota 1: {nota1}")
            if nota2:
                doc.add_paragraph(f"Nota 2: {nota2}")

            # Sección 3: Planilla de Superficies
            self._add_table_superficies(doc)

            # Sección 4: Planilla de Lados
            self._add_table_lados(doc)

            # Sección 5: Croquis y Referencias
            h5 = doc.add_paragraph("5. CROQUIS Y REFERENCIAS")
            h5.runs[0].bold = True
            referencias = self._normalize_refs(self.datos.get("referencias"))
            croquis_text = (self.datos.get("croquis") or "").strip()
            if referencias:
                doc.add_paragraph("Referencias:")
                for ref in referencias:
                    doc.add_paragraph(ref, style="List Bullet")
            if croquis_text:
                doc.add_paragraph(croquis_text)

            # Sección 6: Texto completo (opcional para auditoría)
            texto_completo = self.datos.get("texto_completo")
            if texto_completo:
                h6 = doc.add_paragraph("6. TEXTO COMPLETO (EXTRAÍDO DEL PDF)")
                h6.runs[0].bold = True
                doc.add_paragraph(texto_completo)

            # Sección 7: Coordenadas (si existen)
            self._add_table_coordenadas(doc)

            # Cierre
            doc.add_paragraph("Con esto se dan por finalizadas las operaciones de mensura y división.")
            cierre = doc.add_paragraph(f"Santiago del Estero, {fecha_op}")
            cierre.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            # Pie institucional
            section = doc.sections[0]
            footer = section.footer
            # Limpia footer y agrega texto institucional
            if footer.paragraphs:
                for p in footer.paragraphs:
                    p.text = ""
            footer.add_paragraph("Agrimensores SDE - Santiago del Estero")

            # Guardar
            outputs_dir = os.path.join(settings.MEDIA_ROOT, "outputs", "memorias")
            os.makedirs(outputs_dir, exist_ok=True)
            filename = f"Memoria_{getattr(self.plano, 'id', 'sin_id')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            filepath = os.path.join(outputs_dir, filename)
            doc.save(filepath)

            logger.info("Memoria guardada en: %s", filepath)
            return f"outputs/memorias/{filename}"

        except Exception as e:
            logger.error("Error en DocxGenerator: %s", str(e))
            logger.error("Contexto de datos al fallar: %s", self.datos)
            raise
