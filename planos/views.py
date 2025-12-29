import logging
import os
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib import messages
from django.http import FileResponse, Http404
from django.contrib.auth import authenticate, login
from django.conf import settings
from docx2pdf import convert

from .models import Plano
from .decorators import superuser_required
from .services import procesar_pdf
from django.http import HttpResponse

from django.http import JsonResponse
from planos.utils.pdf_processor import PDFProcessor
# generar_memoria

from planos.utils.ia_memoria import generar_memoria_gemini
from docx import Document

logger = logging.getLogger(__name__)

def login_view(request):
    if request.method == "POST":
        username = request.POST.get("username")
        password = request.POST.get("password")
        user = authenticate(request, username=username, password=password)
        if user is not None:
            login(request, user)
            return redirect("lista_planos")  # redirige al panel
        else:
            messages.error(request, "Usuario o contraseña incorrectos")
    return render(request, "login.html") 

@superuser_required
def upload_plano(request):
    """Vista para subir un plano PDF y procesarlo directamente"""
    if request.method == 'POST':
        titulo = request.POST.get('titulo')
        descripcion = request.POST.get('descripcion', '')
        archivo_pdf = request.FILES.get('archivo_pdf')

        # Validaciones
        if not titulo or not archivo_pdf:
            messages.error(request, 'El título y el archivo PDF son obligatorios.')
            return render(request, 'planos/upload_plano.html', {
                'titulo': titulo,
                'descripcion': descripcion
            })

        if not archivo_pdf.name.lower().endswith('.pdf'):
            messages.error(request, 'El archivo debe ser un PDF.')
            return render(request, 'planos/upload_plano.html')

        if archivo_pdf.size > 10 * 1024 * 1024:
            messages.error(request, 'El archivo no debe superar los 10MB.')
            return render(request, 'planos/upload_plano.html')

        try:
            plano = Plano.objects.create(
                titulo=titulo,
                descripcion=descripcion,
                archivo_pdf=archivo_pdf,
                usuario=request.user if request.user.is_authenticated else None,
                estado='procesando'
            )

            # Procesar PDF y generar memoria
            procesar_pdf(plano)

            messages.success(
                request,
                f'¡Plano "{titulo}" procesado exitosamente! La memoria descriptiva está lista.'
            )
            return redirect('detalle_plano', plano_id=plano.id)

        except Exception as e:
            logger.error(f"Error al cargar plano: {str(e)}")
            messages.error(request, f'Error al cargar el plano: {str(e)}')
            return render(request, 'planos/upload_plano.html')

    return render(request, 'planos/upload_plano.html')

@superuser_required
def lista_planos(request):
    """Vista para listar todos los planos"""
    planos = Plano.objects.all().order_by('-fecha_carga')
    return render(request, 'planos/lista_planos.html', {'planos': planos})

@superuser_required
def detalle_plano(request, plano_id):
    plano = get_object_or_404(Plano, id=plano_id)

    # Como es JSONField, ya viene como dict
    memoria_preview = plano.datos_procesados if plano.datos_procesados else None

    return render(
        request,
        'planos/detalle_plano.html',
        {
            'plano': plano,
            'memoria_preview': memoria_preview
        }
    )

@superuser_required
def descargar_memoria(request, plano_id):
    """Vista para descargar la memoria descriptiva generada en Word"""
    plano = get_object_or_404(Plano, id=plano_id)

    if not plano.memoria_path or plano.estado != 'completado':
        messages.error(request, 'La memoria descriptiva aún no está disponible.')
        return redirect('detalle_plano', plano_id=plano.id)

    memoria_full_path = os.path.join(settings.MEDIA_ROOT, plano.memoria_path)
    if not os.path.exists(memoria_full_path):
        raise Http404("Archivo no encontrado")

    response = FileResponse(
        open(memoria_full_path, 'rb'),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename="Memoria_{plano.titulo}.docx"'
    return response


@superuser_required
def descargar_memoria_pdf(request, plano_id):
    """Vista para convertir la memoria Word a PDF y descargarla"""
    plano = get_object_or_404(Plano, id=plano_id)

    if not plano.memoria_path or plano.estado != 'completado':
        messages.error(request, 'La memoria descriptiva aún no está disponible.')
        return redirect('detalle_plano', plano_id=plano.id)

    memoria_full_path = os.path.join(settings.MEDIA_ROOT, plano.memoria_path)
    pdf_path = memoria_full_path.replace(".docx", ".pdf")

    try:
        convert(memoria_full_path, pdf_path)
    except Exception as e:
        logger.error(f"Error convirtiendo a PDF: {str(e)}")
        messages.error(request, f'Error al convertir a PDF: {str(e)}')
        return redirect('detalle_plano', plano_id=plano.id)

    response = FileResponse(open(pdf_path, 'rb'), content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="Memoria_{plano.titulo}.pdf"'
    return response


@superuser_required
def reprocesar_plano(request, plano_id):
    """Vista para reprocesar un plano"""
    plano = get_object_or_404(Plano, id=plano_id)
    try:
        plano.estado = 'procesando'
        plano.save()
        procesar_pdf(plano)
        messages.success(request, 'Plano reprocesado exitosamente.')
    except Exception as e:
        logger.error(f"Error reprocesando plano {plano_id}: {str(e)}")
        plano.estado = 'error'
        plano.save()
        messages.error(request, f'Error al reprocesar: {str(e)}')
    return redirect('detalle_plano', plano_id=plano.id)


@superuser_required
def subir_memoria_drive(request, plano_id):
    """Vista para subir la memoria descriptiva a Google Drive (simulado)"""
    plano = get_object_or_404(Plano, id=plano_id)

    if not plano.memoria_path or plano.estado != 'completado':
        messages.error(request, 'La memoria descriptiva aún no está disponible.')
        return redirect('detalle_plano', plano_id=plano.id)

    try:
        memoria_full_path = os.path.join(settings.MEDIA_ROOT, plano.memoria_path)
        logger.info(f"Subiendo memoria {memoria_full_path} a Google Drive...")
        # Aquí iría la lógica real de integración con la API de Google Drive
        messages.success(request, 'Memoria subida a Google Drive (simulado).')
    except Exception as e:
        logger.error(f"Error subiendo memoria a Drive: {str(e)}")
        messages.error(request, f'Error al subir a Drive: {str(e)}')

    return redirect('detalle_plano', plano_id=plano.id)

@superuser_required
def ver_reporte(request, plano_id):
    """Vista para ver o descargar el reporte de cumplimiento"""
    plano = get_object_or_404(Plano, id=plano_id)
    # Por ahora devolvemos un texto simple, luego podés generar un Word/PDF igual que la memoria
    return HttpResponse(f"Reporte de cumplimiento para el plano {plano.titulo}")


@superuser_required
def eliminar_plano(request, plano_id):
    """Vista para eliminar un plano"""
    plano = get_object_or_404(Plano, id=plano_id)
    titulo = plano.titulo
    plano.delete()
    messages.success(request, f'Plano "{titulo}" eliminado correctamente.')
    return redirect('lista_planos')

# @superuser_required
# def generar_memoria_preview(request, plano_id):
#     plano = Plano.objects.get(id=plano_id)
#     processor = PDFProcessor(plano.archivo_pdf.path)
#     datos = processor.extract_data()
#     memoria_texto = generar_memoria(datos)
#     return JsonResponse({"memoria": memoria_texto})

@superuser_required
def generar_memoria_preview(request, plano_id):  # noqa: F811
    plano = Plano.objects.get(id=plano_id)
    processor = PDFProcessor(plano.archivo_pdf.path)
    datos = processor.extract_data()

    memoria_texto = generar_memoria_gemini(datos)

    return JsonResponse({"memoria": memoria_texto})

@superuser_required
def descargar_memoria_gemini(request, plano_id):
    plano = Plano.objects.get(id=plano_id)
    processor = PDFProcessor(plano.archivo_pdf.path)
    datos = processor.extract_data()

    # Texto narrativo generado por Gemini
    memoria_texto = generar_memoria_gemini(datos)

    # Crear documento Word
    doc = Document()
    doc.add_heading("Memoria Descriptiva", level=1)
    doc.add_paragraph(memoria_texto)

    # Tabla de superficies
    if datos.get("superficies"):
        doc.add_heading("Planilla de Superficies", level=2)
        table = doc.add_table(rows=1, cols=5)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Lote"
        hdr_cells[1].text = "Sup. s/Título"
        hdr_cells[2].text = "Sup. s/Mensura"
        hdr_cells[3].text = "Diferencia"
        hdr_cells[4].text = "Observaciones"
        for sup in datos["superficies"]:
            row_cells = table.add_row().cells
            row_cells[0].text = sup.get("designacion", "")
            row_cells[1].text = sup.get("sup_titulo", "")
            row_cells[2].text = sup.get("sup_mensura", "")
            row_cells[3].text = sup.get("diferencia", "")
            row_cells[4].text = sup.get("observaciones", "")

    # Tabla de coordenadas
    if datos.get("coordenadas"):
        doc.add_heading("Coordenadas Geodésicas POSGAR 07", level=2)
        table = doc.add_table(rows=1, cols=6)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Punto"
        hdr_cells[1].text = "Latitud"
        hdr_cells[2].text = "Longitud"
        hdr_cells[3].text = "Norte GK"
        hdr_cells[4].text = "Este GK"
        hdr_cells[5].text = "Observación"
        for c in datos["coordenadas"]:
            row_cells = table.add_row().cells
            row_cells[0].text = c.get("punto", "")
            row_cells[1].text = c.get("latitud", "")
            row_cells[2].text = c.get("longitud", "")
            row_cells[3].text = c.get("norte_gk", "")
            row_cells[4].text = c.get("este_gk", "")
            row_cells[5].text = c.get("observacion", "")

    # Notas y referencias
    doc.add_heading("Notas y Referencias", level=2)
    doc.add_paragraph(f"Nota 1: {datos.get('nota1', '')}")
    doc.add_paragraph(f"Nota 2: {datos.get('nota2', '')}")
    doc.add_paragraph(f"Referencias: {datos.get('referencias', '')}")

    # Exportar como DOCX
    response = HttpResponse(content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    response["Content-Disposition"] = f'attachment; filename="memoria_{plano_id}.docx"'
    doc.save(response)
    return response
